import json
import os

from docx import Document
from docx.shared import Cm
from docx.shared import Pt

#LOGIC CODE
class Listing: #Listing object to be used for education and work
    def __init__(self,name,start,end,subtitle,location,details):
        self.name = name
        self.start = start
        self.end = end
        if end == '':
            self.dates = start
        else:
            self.dates = start+" - "+end
        self.subtitle = subtitle
        self.location = location
        self.details = details

#adds instances to above lists
def add_work(workListings, name, start, end, position, location, details):
    workListings.append(Listing(name,start,end,position,location,details))
    return workListings
def add_education(educationListings, name, start, end, years, location, details):
    educationListings.append(Listing(name,start,end,years,location,details))
    return educationListings
def add_skill(skills, skill, details):
    skills.append([skill]+details)
    return skills

#Finds number of tabs needed
def find_tabs(chars):
    if chars < 30:
        return 12
    elif chars > 98:
        return 1
    else:
        return int((0.0021*(chars**2))-((0.3891)*chars)+(21.018))

#Loads data from file and calls functions to store them
def loadData(dir):
    educationListings = []
    workListings = []
    skills = []

    try:
        dataFile = open(dir,'r')
    except FileNotFoundError:
        #dictionary = {'skills': [['Skill goes here', ['detail 1', 'detail 2', 'detail 2']]], 'work': [['Company', 'Start Date', 'End Date', 'Position', 'Location']], 'education': [["Queen's University", '2021', '2025 (expected)', 'BCMPH-I, Computer Science']], 'info': [['Resume', 'John Doe', 'examaple@email.com', '(xxx)xxx-xxxx', 'Location', 'Your description telling your future employer a little about yourself goes here!']]}
        #json_object = json.dumps(dictionary, indent=4)
        #dataFile = open(dir, "x")
        #dataFile.write(json_object)
        skills = add_skill(skills, "Your skill", ["detail 1","detail 2","detail 3"])
        workListings = add_work(workListings, "Company", "Start Date", "End Date", "Position", "Location", ["Detail 1", "Detail 2", "Detail 3"])
        educationListings = add_education(educationListings, "University", "Start Date", "End Date", "Location", "Years", ["Detail 1", "Detail 2", "Detail 3"])
        info = [['Resume', 'John Doe', 'example@email.com', '(xxx)xxx-xxxx', 'Location', 'Your description telling your future employer a little about yourself goes here!']]
        return (workListings, educationListings, skills, info)

    lines = json.load(dataFile)
    dataFile.close()

    for line in lines['skills']:
        skills = add_skill(skills, line[0], line[1])
    for line in lines['work']:
        workListings = add_work(workListings, line[0], line[1], line[2], line[3], line[4], line[5])
    for line in lines['education']:
        educationListings = add_education(educationListings, line[0], line[1], line[2], line[3], line[4], line[5])
    info = lines['info']

    return (workListings, educationListings, skills, info)

#BUILD THE DOCUMENTS
def buildDocuments(workListings, educationListings, skills, info):
    for loop in info:
        document = Document()

        sections = document.sections
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        font = document.styles['Normal'].font
        font.name = "Garamond"

        title = document.add_paragraph()
        title1 = title.add_run(loop[1])
        title1.font.name = "Garamond"
        title1.bold = True
        title1.font.size = Pt(24)

        title2 = title.add_run("\n"+loop[2])
        title2.font.size = Pt(16)

        title3 = title.add_run('  ')
        title3.font.name = "Wingdings"
        title3.font.size = Pt(16)

        title4 = title.add_run(loop[3])
        title4.font.size = Pt(16)
        title4.font.name = "Garamond"

        title5 = title.add_run('  ')
        title5.font.name = "Wingdings"
        title5.font.size = Pt(16)

        title6 = title.add_run(loop[4])
        title6.font.name = "Garamond"
        title6.font.size = Pt(16)

        summary = document.add_paragraph()
        SU = summary.add_run(loop[5])
        SU.font.size = Pt(12)
        
        work_experienceH = document.add_paragraph()
        WE = work_experienceH.add_run("WORK EXPERIENCE")
        WE.font.size = Pt(13)
        WE.bold = True

        for listing in workListings:
            newPara = document.add_paragraph('')
            NP1 = newPara.add_run(listing.name+("\t"*find_tabs(len(listing.name)+len(listing.dates)))+listing.dates)
            NP1.bold = True
            NP1.font.size = Pt(12)

            NP2 = newPara.add_run("\n"+listing.subtitle+("\t"*find_tabs(len(listing.name)+len(listing.location)))+listing.location)
            NP2.bold = False
            NP2.italic = True

            for detail in listing.details:
                NP3 = newPara.add_run("\nq ")
                NP3.italic = False
                NP3.font.name = "Wingdings"
                NP4 = newPara.add_run(detail)
                NP4.font.name = "Garamond"

        educationH = document.add_paragraph()
        E = educationH.add_run("\nEDUCATION")
        E.font.size = Pt(13)
        E.bold = True

        for listing in educationListings:
            newPara = document.add_paragraph('')
            NP1 = newPara.add_run(listing.name+("\t"*find_tabs(len(listing.name)+len(listing.dates)))+listing.dates)
            NP1.bold = True
            NP1.font.size = Pt(12)

            NP2 = newPara.add_run("\n"+listing.subtitle+("\t"*find_tabs(len(listing.subtitle)+len(listing.location)))+listing.location)
            NP2.bold = False
            NP2.italic = True

            for detail in listing.details:
                NP3 = newPara.add_run("\nq ")
                NP3.italic = False
                NP3.font.name = "Wingdings"
                NP4 = newPara.add_run(detail)
                NP4.font.name = "Garamond"

        skillsH = document.add_paragraph()
        S = skillsH.add_run("\nSKILLS, CAPABILITIES, AND OTHER ACHIEVMENTS")
        S.font.size = Pt(13)
        S.bold = True

        NP = document.add_paragraph()
        for skill in skills:
            if skill == skills[0]:
                r = NP.add_run("q ")
            else:
                r = NP.add_run("\nq ")
            r.font.name = "Wingdings"
            r2 = NP.add_run(skill[0]+"; ")
            r2.font.name = "Garamond"
            r2.bold = True
            for s in range(1,len(skill)):
                if s == len(skill)-1:
                    r3 = NP.add_run(skill[s]+"\n")
                else:
                    r3 = NP.add_run(skill[s]+", ")
            

        document.save(loop[0]+".docx")
    print("Built Documents")

if __name__ == '__main__':
    workListings, educationListings, skills, info = loadData(os.getcwd()+"\\data.json")
    buildDocuments(workListings,educationListings,skills,info)
